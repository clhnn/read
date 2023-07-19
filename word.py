import docx
import json

class Word:
    
    #初始化方法，設定要處理的 Word 檔案
    def __init__(self, doc):
        self.doc = docx.Document(doc)

    # 讀取 Word 文件中的所有文本段落
    def read_text(self):
        paragraphs = self.doc.paragraphs
        text_list = [paragraph.text for paragraph in paragraphs]
        return text_list

    # 提取 Word 文件中的所有圖片並儲存成檔案
    def extract_images(self):
        rels = self.doc.part.rels
        images = []     
        for rel in rels:
            rel_type = rels[rel].reltype
            if "image" in rel_type:
                image_part = rels[rel]._target
                image_data = image_part.blob
                image_filename = f"image{len(images)+1}.png"
                with open(image_filename, "wb") as f:
                    f.write(image_data)
                images.append(image_filename)
        return images

    # 讀取 Word 文件中的表格內容並返回列表形式
    def read_table(self):
        tables_data = []
        for table in self.doc.tables:
            data = []
            keys = None
            for i, row in enumerate(table.rows):
                text = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    keys = tuple(text)
                    continue
                row_data = dict(zip(keys, text))
                data.append(row_data)
            tables_data.append(data)
        return tables_data

    # 讀取 Word 文件中的標題樣式段落並返回列表
    def readheading(self):
        paragraphs = self.doc.paragraphs
        heading=[]
        for paragraph in paragraphs:
            if paragraph.style.name.startswith('Heading'):
                heading.append(paragraph.text)
        return heading

#主程式       
if __name__ == '__main__':
    
    # 請將 "your file name.docx" 替換為要處理的實際 Word 檔案名稱
    word = Word("your file name.docx")

    # 讀取 Word 文件中的所有文本段落
    text_list = word.read_text()

    # 提取 Word 文件中的所有圖片並儲存成檔案
    images_list = word.extract_images()

    # 讀取 Word 文件中的表格內容並返回列表形式
    tables_data = word.read_table()

    # 讀取 Word 文件中的標題樣式段落並返回列表
    heading = word.readheading()

    # 將結果整合成字典
    result = {
        'Heading': heading,
        'text': text_list,
        'images': images_list,
        'tables': tables_data        
    }

    # 將結果寫入 JSON 檔案
    with open('wordtxt.json', 'w') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print("JSON result exported to wordtxt.json")    
